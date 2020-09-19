import docx

doc = docx.Document('../resource/word/demo.docx')
print(len(doc.paragraphs))

print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)


print(len(doc.paragraphs[1].runs))

print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)

"""
7
Document Title
A plain paragraph with some bold and some italic
5
A plain paragraph with
 some 
bold
 and some 
"""
