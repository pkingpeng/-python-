import docx

doc = docx.Document('../resource/word/demo.docx')
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)

doc.paragraphs[0].style = 'Normal'

print(doc.paragraphs[1].text)
print(
    doc.paragraphs[1].runs[0].text,
    doc.paragraphs[1].runs[1].text,
    doc.paragraphs[1].runs[2].text,
    doc.paragraphs[1].runs[3].text
)

"""
如果这里 style 设定为 QuoteChar 去掉中间的空格会收到警告，内容如下。
/usr/local/lib/python3.7/site-packages/docx/styles/styles.py:139: UserWarning: style lookup by style_id is deprecated. Use style name as key instead.
  return self._get_style_id_from_style(self[style_name], style_type)
看起来应该是开发团队自己解析对应的格式，而不是强迫用户写出一个必须不带空格的样式。
"""
doc.paragraphs[1].runs[0].style = 'Quote Char'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True

doc.save('../resource/word/output/restyled.docx')
print('Done')
